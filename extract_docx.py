from docx import Document
import json

doc = Document('知识图谱数据(1).docx')

# 提取所有段落
print("=" * 50)
print("文档内容:")
print("=" * 50)

for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"{i}: {para.text}")

# 提取所有表格
print("\n" + "=" * 50)
print("表格内容:")
print("=" * 50)

for table_idx, table in enumerate(doc.tables):
    print(f"\n表格 {table_idx + 1}:")
    for row_idx, row in enumerate(table.rows):
        row_data = [cell.text.strip() for cell in row.cells]
        print(f"  行{row_idx}: {row_data}")
